import json
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

ROOT = Path(__file__).resolve().parent
DATA = json.loads((ROOT / 'content.json').read_text(encoding='utf-8'))
OUT = ROOT.parent / 'KVS_Final_Project.docx'

def set_paragraph_rtl(paragraph, rtl=True):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if rtl else '0')

def set_run_lang(run, lang, rtl=False):
    rPr = run._r.get_or_add_rPr()
    if rtl:
        rtl_el = rPr.find(qn('w:rtl'))
        if rtl_el is None:
            rtl_el = OxmlElement('w:rtl')
            rPr.append(rtl_el)
        rtl_el.set(qn('w:val'), '1')
    lang_el = rPr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rPr.append(lang_el)
    lang_el.set(qn('w:val'), lang)
    if rtl:
        lang_el.set(qn('w:bidi'), lang)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def add_he(doc, text, style='Hebrew Body'):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_rtl(p, True)
    r = p.add_run(text)
    set_run_lang(r, 'he-IL', True)
    return p

def add_ru(doc, text, style='Russian Body'):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_rtl(p, False)
    r = p.add_run(text)
    set_run_lang(r, 'ru-RU', False)
    return p

def add_heading(doc, text, size, rtl=False, align=None):
    p = doc.add_paragraph()
    p.alignment = align if align is not None else (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
    set_paragraph_rtl(p, rtl)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    set_run_lang(r, 'he-IL' if rtl else 'ru-RU', rtl)
    return p

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.3)
sec.right_margin = Cm(2.3)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
styles['Normal'].font.size = Pt(12)
styles['Normal'].paragraph_format.space_after = Pt(7)
styles['Normal'].paragraph_format.line_spacing = 1.25
for name, size, bold in [('Title', 26, True), ('Subtitle', 15, False), ('Heading 1', 19, True), ('Heading 2', 15, True)]:
    st = styles[name]
    st.font.name = 'Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.font.size = Pt(size)
    st.font.bold = bold
for name, size, spacing in [('Hebrew Body', 12, 1.3), ('Russian Body', 11.5, 1.25)]:
    if name not in styles:
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles['Normal']
        st.font.name = 'Arial'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        st.font.size = Pt(size)
        st.paragraph_format.line_spacing = spacing
        st.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph_rtl(p, True)
r = p.add_run(DATA['hebrew_title']); r.bold = True; r.font.size = Pt(26); set_run_lang(r, 'he-IL', True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph_rtl(p, True)
r = p.add_run(DATA['hebrew_subtitle']); r.bold = True; r.font.size = Pt(18); set_run_lang(r, 'he-IL', True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph_rtl(p, True)
r = p.add_run(DATA['hebrew_kind']); r.font.size = Pt(15); set_run_lang(r, 'he-IL', True)
for _ in range(7): doc.add_paragraph()
info = doc.add_table(rows=3, cols=2); info.alignment = 1; info.autofit = False
rows = [('שם התלמיד:', DATA['student']), ('שם הקורס:', DATA['course']), ('מצב המסמך:', DATA['status'])]
for row, (label, value) in zip(info.rows, rows):
    row.cells[0].width = Cm(5); row.cells[1].width = Cm(8)
    for cell in row.cells: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p0=row.cells[0].paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_paragraph_rtl(p0, True)
    rr=p0.add_run(value); set_run_lang(rr,'he-IL',True)
    p1=row.cells[1].paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_paragraph_rtl(p1, True)
    rr=p1.add_run(label); rr.bold=True; set_run_lang(rr,'he-IL',True); set_cell_shading(row.cells[1],'E9EEF6')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(DATA['year']).font.size=Pt(12)

doc.add_page_break()
add_heading(doc, DATA['hebrew_note_title'], 15, rtl=True)
add_he(doc, DATA['hebrew_note'])
add_heading(doc, DATA['russian_note_title'], 15, rtl=False)
add_ru(doc, DATA['russian_note'])
doc.add_paragraph()
add_heading(doc, 'תוכן עניינים זמני', 17, rtl=True)
for line in DATA['toc']:
    p=doc.add_paragraph(style='Hebrew Body'); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_paragraph_rtl(p,True)
    rr=p.add_run('• '+line); set_run_lang(rr,'he-IL',True)

doc.add_page_break()
add_heading(doc, DATA['chapter_he_title'], 20, rtl=True)
for text in DATA['chapter_he_paragraphs']: add_he(doc, text)

doc.add_page_break()
add_heading(doc, 'Рабочий перевод на русский язык', 14, rtl=False)
add_heading(doc, DATA['chapter_ru_title'], 19, rtl=False)
for text in DATA['chapter_ru_paragraphs']: add_ru(doc, text)
add_heading(doc, 'Источники и библиография', 15, rtl=False)
add_ru(doc, DATA['sources_note'])

props=doc.core_properties
props.title='KVS Final Project — Second Temple Groups and Hellenistic vs Jewish Identity'
props.subject='Historical introduction and play, bilingual working draft'
props.author='Arseny Perel'
props.keywords='KVS, Second Temple, Hellenism, Jewish identity, Pharisees, Sadducees, Essenes, Zealots, Nazarenes'
doc.save(OUT)
print(OUT)
